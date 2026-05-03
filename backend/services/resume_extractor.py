"""
resume_extractor.py
Handles multi-format file parsing: PDF, DOCX, DOC, JPG, PNG, TXT, RTF
"""

import os
import io
import tempfile
from pathlib import Path
from fastapi import UploadFile

# ── PDF ──────────────────────────────────────────────────────────────────────
def extract_from_pdf(file_bytes: bytes) -> str:
    import pdfplumber
    text = []
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text.append(page_text)
    return "\n".join(text)

# ── DOCX ─────────────────────────────────────────────────────────────────────
def extract_from_docx(file_bytes: bytes) -> str:
    from docx import Document
    doc = Document(io.BytesIO(file_bytes))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    paragraphs.append(cell.text.strip())
    return "\n".join(paragraphs)

# ── IMAGE (OCR) ───────────────────────────────────────────────────────────────
def extract_from_image(file_bytes: bytes) -> str:
    try:
        import easyocr
        import numpy as np
        from PIL import Image

        img = Image.open(io.BytesIO(file_bytes)).convert("RGB")
        img_array = np.array(img)
        reader = easyocr.Reader(["en"], gpu=False)
        results = reader.readtext(img_array)
        return " ".join([text for (_, text, _) in results])
    except Exception:
        # Fallback to pytesseract
        import pytesseract
        from PIL import Image
        img = Image.open(io.BytesIO(file_bytes))
        return pytesseract.image_to_string(img)

# ── RTF ───────────────────────────────────────────────────────────────────────
def extract_from_rtf(file_bytes: bytes) -> str:
    from striprtf.striprtf import rtf_to_text
    return rtf_to_text(file_bytes.decode("utf-8", errors="ignore"))

# ── TXT ───────────────────────────────────────────────────────────────────────
def extract_from_txt(file_bytes: bytes) -> str:
    return file_bytes.decode("utf-8", errors="ignore")

# ── ROUTER ────────────────────────────────────────────────────────────────────
async def extract_text_from_file(file: UploadFile) -> dict:
    """
    Auto-detects file type and extracts raw text from any resume format.
    Returns { text, format_detected, success, error }
    """
    file_bytes = await file.read()
    filename = file.filename.lower()
    ext = Path(filename).suffix

    format_map = {
        ".pdf": ("PDF", extract_from_pdf),
        ".docx": ("DOCX", extract_from_docx),
        ".doc": ("DOC (legacy)", extract_from_docx),
        ".txt": ("Plain Text", extract_from_txt),
        ".rtf": ("RTF", extract_from_rtf),
        ".jpg": ("JPEG Image (OCR)", extract_from_image),
        ".jpeg": ("JPEG Image (OCR)", extract_from_image),
        ".png": ("PNG Image (OCR)", extract_from_image),
        ".webp": ("WebP Image (OCR)", extract_from_image),
        ".bmp": ("BMP Image (OCR)", extract_from_image),
        ".tiff": ("TIFF Image (OCR)", extract_from_image),
        ".tif": ("TIFF Image (OCR)", extract_from_image),
    }

    if ext not in format_map:
        return {
            "text": "",
            "format_detected": "Unknown",
            "success": False,
            "error": f"Unsupported file format: {ext}. Please upload PDF, DOCX, DOC, TXT, RTF, JPG, PNG, WEBP, BMP, or TIFF."
        }

    format_label, extractor = format_map[ext]
    try:
        raw_text = extractor(file_bytes)
        if not raw_text or len(raw_text.strip()) < 50:
            return {
                "text": raw_text,
                "format_detected": format_label,
                "success": False,
                "error": "Could not extract enough text. The file may be empty, password-protected, or a scanned image with no readable text."
            }
        return {
            "text": raw_text.strip(),
            "format_detected": format_label,
            "success": True,
            "error": None
        }
    except Exception as e:
        return {
            "text": "",
            "format_detected": format_label,
            "success": False,
            "error": f"Extraction failed: {str(e)}"
        }
